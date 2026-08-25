# python-docx ships no type stubs; relax strict checks for its surface.
# pyright: reportUnknownMemberType=false, reportUnknownParameterType=false
# pyright: reportUnknownVariableType=false, reportMissingTypeStubs=false, reportGeneralTypeIssues=false
from io import BytesIO

from docx import Document

from worker.documents.renderer import render_resume_docx


def paragraphs(document: Document) -> list[str]:
	return [paragraph.text for paragraph in document.paragraphs]


def facts() -> dict[str, object]:
	return {
		"contact": {
			"name": "Ada Lovelace",
			"email": "ada@example.com",
			"phone": None,
			"location": "London",
		},
		"skills": [
			{"canonicalName": "Python", "category": "Data science"},
			{"canonicalName": "PostgreSQL", "category": "Databases"},
			{"canonicalName": "Kafka", "category": None},
		],
		"employment": [
			{
				"title": "Engineer",
				"employer": "Example Corp",
				"startDate": "2022-03",
				"endDate": None,
				"isCurrent": True,
			}
		],
		"education": [{"degree": "B.Tech", "institution": "VJTI"}],
		"certifications": [],
	}


def suggestions() -> list[dict[str, object]]:
	return [{"title": "Quantify outcomes", "detail": "Add measured results."}]


def test_rendered_docx_contains_documented_sections() -> None:
	document = Document(BytesIO(render_resume_docx(facts(), suggestions())))
	texts = paragraphs(document)
	assert texts[0] == "Ada Lovelace"
	assert any("ada@example.com · London" in line for line in texts)
	assert any(line.startswith("Data science: Python") for line in texts)
	assert any(line.startswith("Other: Kafka") for line in texts)
	assert any("Engineer — Example Corp (2022-03 to present)" in line for line in texts)
	assert any("B.Tech, VJTI" in line for line in texts)
	assert any("Improvement notes" == line for line in texts)
	assert any("Quantify outcomes: Add measured results." in line for line in texts)


def test_rendered_docx_omits_empty_sections_and_never_invents_content() -> None:
	minimal: dict[str, object] = {"contact": {"name": None}, "skills": []}
	document = Document(BytesIO(render_resume_docx(minimal, [])))
	texts = paragraphs(document)
	assert texts[0] == "Resume"
	headings = {"Skills", "Experience", "Education", "Certifications", "Improvement notes"}
	assert not any(line in headings for line in texts)


def test_rendered_docx_strips_control_characters_from_facts() -> None:
	unsanitized: dict[str, object] = {
		"contact": {"name": "Ada\x0bLovelace\x00", "email": "ada@example.com"}
	}
	document = Document(BytesIO(render_resume_docx(unsanitized, [])))
	assert document.paragraphs[0].text == "AdaLovelace"


def test_rendered_docx_caps_individual_fact_length() -> None:
	long_name = "A" * 5_000
	fact: dict[str, object] = {"contact": {"name": long_name}}
	document = Document(BytesIO(render_resume_docx(fact, [])))
	assert len(document.paragraphs[0].text) == 2_000
