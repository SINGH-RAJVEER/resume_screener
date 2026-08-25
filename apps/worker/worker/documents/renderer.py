from collections.abc import Mapping, Sequence
from io import BytesIO
from typing import cast

from docx import Document


def render_resume_docx(
	facts: Mapping[str, object], suggestions: Sequence[Mapping[str, object]]
) -> bytes:
	# The corrected document only reorganizes documented facts; it never adds
	# content the evaluation did not find in the source resume.
	document = Document()
	contact = mapping_of(facts.get("contact"))
	name = text_of(contact.get("name")) or "Resume"
	document.add_heading(name, level=0)
	contact_line = " · ".join(
		item
		for item in (
			text_of(contact.get("email")),
			text_of(contact.get("phone")),
			text_of(contact.get("location")),
		)
		if item
	)
	if contact_line:
		document.add_paragraph(contact_line)

	skills = entries_of(facts.get("skills"))
	if skills:
		document.add_heading("Skills", level=1)
		grouped: dict[str, list[str]] = {}
		for skill in skills:
			canonical = text_of(skill.get("canonicalName"))
			if not canonical:
				continue
			category = text_of(skill.get("category")) or "Other"
			grouped.setdefault(category, []).append(canonical)
		for category in sorted(grouped):
			document.add_paragraph(
				f"{category}: {', '.join(sorted(grouped[category], key=str.casefold))}"
			)

	employment = entries_of(facts.get("employment"))
	if employment:
		document.add_heading("Experience", level=1)
		for role in employment:
			title = " — ".join(
				item
				for item in (
					text_of(role.get("title")),
					text_of(role.get("employer")),
				)
				if item
			)
			start = text_of(role.get("startDate"))
			end = text_of(role.get("endDate"))
			if role.get("isCurrent") is True and start:
				dates = f"{start} to present"
			elif start and end:
				dates = f"{start} to {end}"
			else:
				dates = start or end or ""
			line = f"{title} ({dates})" if title and dates else title or dates or "Role"
			document.add_paragraph(line)

	education = entries_of(facts.get("education"))
	if education:
		document.add_heading("Education", level=1)
		for entry in education:
			parts = [
				item
				for item in (
					text_of(entry.get("degree")),
					text_of(entry.get("fieldOfStudy")),
					text_of(entry.get("institution")),
				)
				if item
			]
			if parts:
				document.add_paragraph(", ".join(parts))

	certifications = entries_of(facts.get("certifications"))
	if certifications:
		document.add_heading("Certifications", level=1)
		for entry in certifications:
			parts = [
				item
				for item in (
					text_of(entry.get("name")),
					text_of(entry.get("issuer")),
				)
				if item
			]
			if parts:
				document.add_paragraph(" — ".join(parts))

	if suggestions:
		document.add_heading("Improvement notes", level=1)
		for suggestion in suggestions:
			title = text_of(suggestion.get("title"))
			detail = text_of(suggestion.get("detail"))
			line = f"{title}: {detail}" if title and detail else title or detail
			if line:
				document.add_paragraph(line, style="List Bullet")

	buffer = BytesIO()
	document.save(buffer)
	return buffer.getvalue()


def mapping_of(value: object) -> Mapping[str, object]:
	if isinstance(value, Mapping):
		return cast(Mapping[str, object], value)
	return {}


def entries_of(value: object) -> list[Mapping[str, object]]:
	if not isinstance(value, list):
		return []
	return [
		cast(Mapping[str, object], item)
		for item in cast(list[object], value)
		if isinstance(item, Mapping)
	]


def text_of(value: object) -> str | None:
	return value if isinstance(value, str) and value.strip() else None
